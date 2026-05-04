#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
中建四局公文格式化核心逻辑
从 app.py 抽取，供 Flask 后端调用
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from constants import *
from detect import (detect_level, is_main_title, HeadingCounter,
                    apply_heading_format, clean_text, set_run_font,
                    set_para_spacing, set_para_indent, resolve_final_level)
from check import (check_punctuation_issues, check_subheading_issues,
                   check_h3_numbering_issues, check_word_numbering_format,
                   check_missing_h2, check_title_punctuation, check_title_trailing_punct,
                   check_numbering_separator,
                   check_list_numbering_restart)
from comment import _apply_comments_to_doc, _add_page_number
from table import extract_table_data, rebuild_table

# Pt/Cm 包装：constants 中存裸值，此处转为 docx 对象
SIZE_CHUHAO  = Pt(SIZE_CHUHAO)
SIZE_ERHAO   = Pt(SIZE_ERHAO)
SIZE_SANHAO  = Pt(SIZE_SANHAO)
SIZE_XIAOSI  = Pt(SIZE_XIAOSI)
MARGIN_TOP    = Cm(MARGIN_TOP)
MARGIN_BOTTOM = Cm(MARGIN_BOTTOM)
MARGIN_LEFT   = Cm(MARGIN_LEFT)
MARGIN_RIGHT  = Cm(MARGIN_RIGHT)


def format_document(src_path: str, dst_path: str):
    """主转换函数：读取 → 应用公文格式 → 保存。
    
    返回格式：(dst_path, warnings_list)
    warnings_list 中每项为 dict，包含 type 和 detail 字段。
    """
    ext = Path(src_path).suffix.lower()
    paragraphs_text = []
    warnings = []
    bold_runs_by_elem = {}  # 函数级变量，跨分支共享

    if ext == '.docx':
        src_doc = Document(src_path)
        body = src_doc.element.body
        ordered_elements = []
        for child in body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'p':
                ordered_elements.append(('p', child))
            elif tag == 'tbl':
                ordered_elements.append(('tbl', child))

        # 预建逐 run 加粗映射
        bold_runs_by_elem.clear()
        for p_obj in src_doc.paragraphs:
            runs_info = [(len(r.text), bool(r.font.bold)) for r in p_obj.runs if r.text and r.text.strip()]
            if runs_info:
                bold_runs_by_elem[p_obj._element] = runs_info

        abstract_num_defs = {}
        num_to_abstract = {}
        try:
            numbering_part = src_doc.part.numbering_part
            if numbering_part:
                num_root = numbering_part._element
                for abstract_num in num_root.findall(qn('w:abstractNum')):
                    an_id = abstract_num.get(qn('w:abstractNumId'))
                    levels = {}
                    for lvl in abstract_num.findall(qn('w:lvl')):
                        ilvl_val = lvl.get(qn('w:ilvl'), '0')
                        num_fmt = lvl.find(qn('w:numFmt'))
                        lvl_text = lvl.find(qn('w:lvlText'))
                        if num_fmt is not None and lvl_text is not None:
                            levels[ilvl_val] = (
                                num_fmt.get(qn('w:val'), ''),
                                lvl_text.get(qn('w:val'), '')
                            )
                    abstract_num_defs[an_id] = levels
                for num in num_root.findall(qn('w:num')):
                    numId = num.get(qn('w:numId'))
                    aRef = num.find(qn('w:abstractNumId'))
                    if aRef is not None:
                        num_to_abstract[numId] = aRef.get(qn('w:val'))
        except Exception:
            pass

        for etype, elem in ordered_elements:
            if etype == 'p':
                text = ''.join(t.text for t in elem.iter(qn('w:t')) if t.text)
                bold_runs = bold_runs_by_elem.get(elem, [])
                is_bold = any(b for _, b in bold_runs) if bold_runs else False
                num_id = num_ilvl = None
                pPr = elem.find(qn('w:pPr'))
                if pPr is not None:
                    numPr = pPr.find(qn('w:numPr'))
                    if numPr is not None:
                        nid_el = numPr.find(qn('w:numId'))
                        ilvl_el = numPr.find(qn('w:ilvl'))
                        if nid_el is not None:
                            num_id = nid_el.get(qn('w:val'))
                        if ilvl_el is not None:
                            num_ilvl = ilvl_el.get(qn('w:val'))
                word_num_level = 'PENDING' if (num_id and num_id != '0') else None
                paragraphs_text.append(('p', text, is_bold, word_num_level, num_id, num_ilvl, bold_runs))
            elif etype == 'tbl':
                result = extract_table_data(elem)
                if result:
                    paragraphs_text.append(result)

        # ──── 预合并：合并被Word拆分的标题碎片 ────
        # 临时禁用：此逻辑导致标题段落被错误合并，编号错乱
        # 直接跳过预合并，使用原始段落列表
        pass  # 下面的while循环已被禁用
        while False and i < len(paragraphs_text):
            if merged[i]:
                i += 1
                continue
            item = paragraphs_text[i]
            if item[0] != 'p':
                i += 1
                continue
            text = item[1]
            # 碎片判断：短文本+无句内标点+无编号前缀+非纯数字+非问候语+不以逗号/顿号/冒号结尾
            # 重要：标题末尾的冒号不应被合并（如"（一）项目上线与业财集成："）
            greeting_kw = '领导|同事|各位|尊敬|您好|下午好|上午好|你好'
            is_greeting_text = (
                re.match(r'^.{2,30}[：:]$', text.strip())
                and re.search(greeting_kw, text.strip())
            )
            # 排除以逗号/顿号/冒号结尾的碎片
            # - 逗号/顿号结尾：句子中间的分隔，不是标题续行
            # - 冒号结尾：标题（如"（一）xxx："），不是正文碎片
            ends_with_sep = bool(re.search(r'[，、：:]\s*$', text))
            is_fragment = (
                1 < len(text) <= 25
                and not re.search(r'[。；！？]', text)
                and not ends_with_sep
                and not re.match(r'^[一二三四五六七八九十]+、', text)
                and not re.match(r'^\d+[.、．]', text)
                and not re.match(r'^（\d+）', text)
                and not re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', text)
                and not re.match(r'^[\d,.\-+%：:（）()]+$', text)
                and not is_greeting_text
            )
            if is_fragment and i > 0 and not merged[i - 1]:
                prev = paragraphs_text[i - 1]
                prev_text = prev[1] if prev[0] == 'p' else ''
                # 跳过空段落，向前继续找
                if not prev_text.strip():
                    i += 1
                    continue
                # 前段也是短碎片（≤25字）且不以标点结尾 → 合并
                # 但如果前段明显较长（>20字），不合并
                prev_is_fragment = (prev[0] == 'p' and 1 < len(prev_text) <= 25
                                    and not re.search(r'[。；！？]', prev_text))
                # 保护：如果前段明显较长（>20字），不合并
                prev_is_long_fragment = (prev[0] == 'p' and len(prev_text) > 20)
                # 临时修复：禁用merge_to_main_title，防止标题段落被错误合并
                merge_to_main_title = False
                # 原代码（暂时禁用）：
                # prev_is_bold_main = (prev[0] == 'p' and prev[2] is True)
                # prev_is_main = (prev[0] == 'p' and is_main_title(prev_text))
                # merge_to_main_title = (prev_is_bold_main or prev_is_main) and len(text) <= 15
                # 禁止合并：当前段或前段看起来像独立条目（含冒号的数据行、月底节点等）
                # 关键修复：禁止合并包含编号前缀的段落（如"1." "2." "（一）"等）
                has_num_prefix_current = bool(
                    re.match(r'^[一二三四五六七八九十]+、', text)
                    or re.match(r'^（[一二三四五六七八九十]+）', text)
                    or re.match(r'^\d+[.、．]\s*', text)
                    or re.match(r'^（\d+）', text)
                    or re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', text)
                )
                has_num_prefix_prev = bool(
                    re.match(r'^[一二三四五六七八九十]+、', prev_text)
                    or re.match(r'^（[一二三四五六七八九十]+）', prev_text)
                    or re.match(r'^\d+[.、．]\s*', prev_text)
                    or re.match(r'^（\d+）', prev_text)
                    or re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', prev_text)
                )
                looks_like_list_item = (
                    bool(re.search(r'[：:]', text))
                    or bool(re.search(r'[：:]', prev_text))
                    or has_num_prefix_current
                    or has_num_prefix_prev
                )
                if (prev_is_fragment and not prev_is_long_fragment and not looks_like_list_item) or merge_to_main_title:
                    # 合并到前一段
                    merged_text = prev_text + text
                    paragraphs_text[i - 1] = (prev[0], merged_text, prev[2], prev[3], prev[4], prev[5])
                    merged[i] = True
                    i += 1
                    continue
            i += 1
        # 预合并已被禁用，直接设置merged为全False
        merged = [False] * len(paragraphs_text)
        # 过滤掉已合并的碎片
        paragraphs_text = [item for i, item in enumerate(paragraphs_text) if not merged[i]]

    elif ext in ('.txt', '.md'):
        num_to_abstract = {}
        abstract_num_defs = {}
        with open(src_path, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f.read().splitlines():
                paragraphs_text.append(('p', line, False, None, None, None))
    else:
        raise ValueError(f"不支持的文件格式：{ext}（仅支持 .docx .txt .md）")

    # 建立 numId → ilvl → 公文层级映射
    numid_ilvl_level_map = {}
    used_num_ids = {
        item[4] for item in paragraphs_text
        if item[0] == 'p' and len(item) > 4 and item[4] and item[4] != '0'
    }
    sorted_num_ids = sorted(used_num_ids, key=lambda x: int(x) if x.isdigit() else 0)

    for nid in sorted_num_ids:
        an_id = num_to_abstract.get(nid)
        if an_id is None:
            continue
        levels = abstract_num_defs.get(an_id, {})
        for ilvl, (fmt, txt) in levels.items():
            # decimal 格式（1. 2. 3.）通常是正文编号列表，不应映射为标题层级
            if fmt == 'decimal':
                continue
            if fmt in ('chineseCounting', 'chineseCountingThousand',
                       'upperLetter', 'lowerLetter',
                       'ideographDigital', 'ideographEnclosedCircle'):
                mapping = {'0': 'h1', '1': 'h2', '2': 'h3'}
                if ilvl in mapping:
                    numid_ilvl_level_map[(nid, ilvl)] = mapping[ilvl]
            else:
                if txt.startswith('（') or txt.startswith('('):
                    mapping = {'0': 'h2', '1': 'h3', '2': 'h4'}
                elif txt.endswith('.') or txt.endswith('、') or txt.endswith('．'):
                    mapping = {'0': 'h3', '1': 'h4'}
                else:
                    mapping = {'0': 'h1'}
                if ilvl in mapping:
                    numid_ilvl_level_map[(nid, ilvl)] = mapping[ilvl]

    if not numid_ilvl_level_map:
        # 仅当 numId 对应的格式不是 decimal 时才创建回退映射
        # decimal 格式的编号列表通常是正文，不应映射为标题
        for order, nid in enumerate(sorted_num_ids):
            # 检查该 numId 对应的格式是否为 decimal
            an_id = num_to_abstract.get(nid)
            is_decimal_only = False
            if an_id:
                levels = abstract_num_defs.get(an_id, {})
                fmts = [fmt for fmt, _ in levels.values()]
                if fmts and all(f == 'decimal' for f in fmts):
                    is_decimal_only = True
            if is_decimal_only:
                continue  # decimal 格式不映射为标题
            level_map = {0: 'h1', 1: 'h2', 2: 'h3'}
            if order in level_map:
                numid_ilvl_level_map[(nid, '0')] = level_map[order]

    # 回填 PENDING
    # 注意：对于十进制编号（numId对应decimal格式）的段落，应谨慎识别为标题
    # Word的十进制编号常用于正文中的编号列表，不应轻易识别为标题层级
    for idx, item in enumerate(paragraphs_text):
        if item[0] == 'p' and item[3] == 'PENDING':
            nid, nilvl = item[4], item[5] or '0'
            resolved_level = numid_ilvl_level_map.get((nid, nilvl))
            text_content = item[1] if len(item) > 1 else ''
            
            # 如果numId对应的是十进制编号（decimal），应谨慎处理
            # 只有当文本看起来像标题（短且无句末标点）时才识别为标题
            # 否则应识别为正文
            an_id = num_to_abstract.get(nid)
            if an_id:
                levels = abstract_num_defs.get(an_id, {})
                fmt, _ = levels.get(nilvl, (None, None))
                if fmt == 'decimal':
                    # 判断文本是否像标题（短、无句末标点、无数字前缀）
                    looks_like_title = (
                        len(text_content) <= 30
                        and not re.search(r'[。；！？]', text_content)
                        and not re.match(r'^\d+[.、．]', text_content)
                    )
                    # 即使看起来像标题，也要验证文本是否有对应格式的编号前缀
                    # 例如：chineseCounting 格式应有"一、"等前缀，decimal 格式应有"1."等前缀
                    if looks_like_title:
                        # 检查文本是否有对应格式的编号前缀
                        has_correct_prefix = False
                        if fmt == 'chineseCounting':
                            # 检查是否有中文数字前缀（一、 二、 三、）
                            has_correct_prefix = bool(re.match(r'^[一二三四五六七八九十]+[、）]', text_content))
                        elif fmt == 'decimal':
                            # 检查是否有阿拉伯数字前缀（1. 2. 3.）
                            has_correct_prefix = bool(re.match(r'^\d+[.、．]', text_content))
                        
                        if not has_correct_prefix:
                            resolved_level = None  # 无编号前缀，不是标题
                    else:
                        resolved_level = None  # 不像标题，肯定是正文
            
            paragraphs_text[idx] = (
                item[0], item[1], item[2],
                resolved_level,
                item[4], item[5]
            )

    # ──── 计算 Word 编号的实际序号（编号列表检测需要） ────
    num_seq = {}
    num_seq_count = {}
    for _pi, _pitem in enumerate(paragraphs_text):
        if _pitem[0] != 'p':
            continue
        _nid = _pitem[4] if len(_pitem) > 4 else None
        _nilvl = _pitem[5] if len(_pitem) > 5 else None
        if _nid and _nid != '0':
            _key = (_nid, _nilvl)
            num_seq_count[_key] = num_seq_count.get(_key, 0) + 1
            num_seq[_pi] = num_seq_count[_key]

    # 运行审查检测
    punct_issues = check_punctuation_issues(paragraphs_text)
    if punct_issues:
        warnings.append({
            'type': '句末标点缺失',
            'detail': f'共 {len(punct_issues)} 处段落可能缺少句末标点（。）：\n'
                      + '\n'.join(f'  段落{idx+1}: "{txt}…"' for idx, txt in punct_issues)
        })

    subhead_issues = check_subheading_issues(paragraphs_text)
    if subhead_issues:
        detail_lines = []
        for idx, txt, major, minor, content in subhead_issues:
            detail_lines.append(f'  段落{idx+1}: "{txt}" — 原文使用 {major}.{minor} 多级编号')
        detail_lines.append('  建议：可将多级编号改为四级标题①②③格式，请人工确认。')
        warnings.append({
            'type': '子标题序号格式',
            'detail': f'检测到 {len(subhead_issues)} 处多级编号（X.Y格式）：\n' + '\n'.join(detail_lines)
        })

    h3_issues = check_h3_numbering_issues(paragraphs_text)
    if h3_issues:
        detail_lines = []
        for idx, txt, num, word in h3_issues:
            detail_lines.append(f'  段落{idx+1}: "{txt}" — 编号 "{num}." 后直接跟"{word}"')
        detail_lines.append('  建议：此类编号建议改为 "一是…""二是…" 格式，请人工确认。')
        warnings.append({
            'type': '三级编号不规范',
            'detail': f'检测到 {len(h3_issues)} 处编号后直接跟动词：\n' + '\n'.join(detail_lines)
        })

    # 检查Word自动编号格式（阿拉伯数字1.2.3.）— Bug 1修复
    word_num_issues = check_word_numbering_format(paragraphs_text, num_to_abstract, abstract_num_defs)
    word_num_indices = {i for i, _ in word_num_issues}
    
    # 检查 Word 编号分隔符（、．→ .）
    sep_issues = check_numbering_separator(paragraphs_text, num_to_abstract, abstract_num_defs)
    sep_indices = {i for i, _, _, _ in sep_issues}
    if sep_issues:
        detail_lines = []
        for idx, txt, bad, good in sep_issues:
            detail_lines.append(f'  段落{idx+1}: "{txt}" — 编号分隔符为"{bad}"，建议改为"{good}"')
        warnings.append({
            'type': '编号分隔符不规范',
            'detail': f'检测到 {len(sep_issues)} 处编号后使用顿号或全角点，建议改为点号"."：\n' + '\n'.join(detail_lines)
        })
    
    # 检查一级标题下直接使用三级标题的情况
    missing_h2_issues = check_missing_h2(paragraphs_text)
    missing_h2_indices = {i for i, _ in missing_h2_issues}  # 用于批注
    if word_num_issues:
        detail_lines = []
        for idx, txt in word_num_issues:
            detail_lines.append(f'  段落{idx+1}: "{txt}" — 使用Word自动编号（1. 2. 3.）')
        detail_lines.append('  建议：可改为二级标题（（一）（二））格式，请人工确认。')
        warnings.append({
            'type': 'Word自动编号格式',
            'detail': f'检测到 {len(word_num_issues)} 处Word自动编号（阿拉伯数字格式）：\n' + '\n'.join(detail_lines)
        })
    
    # 检查一级标题下直接使用三级标题的情况
    if missing_h2_issues:
        detail_lines = []
        for idx, txt in missing_h2_issues:
            detail_lines.append(f'  段落{idx+1}: "{txt}" — 一级标题下直接使用三级标题')
        detail_lines.append('  建议：在一级标题和三级标题之间补充二级标题（一）（二）')
        warnings.append({
            'type': '缺少二级标题',
            'detail': f'检测到 {len(missing_h2_issues)} 处一级标题下直接使用三级标题：\n' + '\n'.join(detail_lines)
        })

    # 检查标题编号标点是否符合规范
    title_punct_check_results = check_title_punctuation(paragraphs_text)
    title_punct_indices = {i for i, _, _, _ in title_punct_check_results}
    
    if title_punct_check_results:
        detail_lines = []
        for idx, txt, issue_type, suggestion in title_punct_check_results:
            detail_lines.append(f'  段落{idx+1}: "{txt}" — {suggestion}')
        warnings.append({
            'type': '标题编号标点不规范',
            'detail': f'检测到 {len(title_punct_check_results)} 处标题编号标点不符合规范：\n' + '\n'.join(detail_lines)
        })
    
    # ──── 标题句末标点检测 ────
    title_punct_issues = check_title_trailing_punct(paragraphs_text, num_to_abstract, abstract_num_defs)

    # 建立审查高亮标记集合
    punct_para_indices = {idx for idx, _ in punct_issues}          # 句末标点缺失
    subhead_para_indices = {idx for idx, _, _, _, _ in subhead_issues}  # X.Y多级编号
    h3_para_indices = {idx for idx, _, _, _ in h3_issues}          # X.是编号
    title_punct_para_indices = {idx for idx, _, _ in title_punct_issues}  # 标题句末标点

    # 标题句末标点检测
    if title_punct_issues:
        detail_lines = []
        for idx, txt, punct in title_punct_issues:
            detail_lines.append(f'  段落{idx+1}: "{txt}" — 标题末尾不应有"{punct}"')
        warnings.append({
            'type': '标题句末标点',
            'detail': f'检测到 {len(title_punct_issues)} 处标题包含句末标点（标题末尾不应有标点符号）：\n'
                      + '\n'.join(detail_lines)
        })

    # 检查一级标题下编号列表的起始编号
    list_num_issues = check_list_numbering_restart(paragraphs_text, num_seq)
    if list_num_issues:
        detail_lines = []
        for prefix, suggestion, _ in list_num_issues:
            detail_lines.append(f'  {suggestion}')
        warnings.append({
            'type': '编号序号错误',
            'detail': f'检测到 {len(list_num_issues)} 处编号序号错误：\n' + '\n'.join(detail_lines)
        })

    # 新建文档
    doc = Document()
    comment_list = []  # 收集批注: [(匹配文本前缀, comment_text), ...]  # 收集批注: [(匹配文本前缀, comment_text), ...]
    section = doc.sections[0]
    section = doc.sections[0]
    section.top_margin    = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin   = MARGIN_LEFT
    section.right_margin  = MARGIN_RIGHT
    section = doc.sections[0]
    section.top_margin    = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin   = MARGIN_LEFT
    section.right_margin  = MARGIN_RIGHT

    normal_style = doc.styles['Normal']
    normal_style.font.name = FONT_TIMES_NEW_ROMAN
    normal_style.font.size = SIZE_SANHAO
    nf = normal_style.element.find('.//' + qn('w:rFonts'))
    if nf is None:
        rPr = normal_style.element.find('.//' + qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            normal_style.element.append(rPr)
        nf = OxmlElement('w:rFonts')
        rPr.insert(0, nf)
    nf.set(qn('w:eastAsia'), FONT_FANGSONG)

    # ──── 预扫描：检测文档是否用 X. 作为顶层编号（无中文一、二、三、） ────
    has_cn_h1 = False
    for item in paragraphs_text:
        if item[0] == 'p':
            t = item[1].strip()
            if re.match(r'^[一二三四五六七八九十]+、', t):
                has_cn_h1 = True
                break
    # 如果没有中文一级编号，但存在 X、或 X.（非X.Y）编号段落，
    # 且 X. 后面的子标题是 （X）或 (X)，则提升所有 X. 为 h1
    promote_x_to_h1 = False
    promote_body_indices = set()  # 需要提升为 h1 的 body 级段落索引
    if not has_cn_h1:
        x_prefix_paras = []  # X. 或 X、 开头的段落索引
        for idx, item in enumerate(paragraphs_text):
            if item[0] == 'p':
                t = item[1].strip()
                if re.match(r'^\d+[.、．]\s*\S', t) and not re.match(r'^\d+\.\d+', t):
                    x_prefix_paras.append(idx)
        if x_prefix_paras:
            # 检查任意一个 X. 段落后是否有 （X）或 (X) 子标题
            has_sub_level = False
            for xpi in x_prefix_paras[:5]:
                for j in range(xpi + 1, min(xpi + 6, len(paragraphs_text))):
                    if paragraphs_text[j][0] == 'p':
                        sub_t = paragraphs_text[j][1].strip()
                        if re.match(r'^（\d+）', sub_t) or re.match(r'^[（(]\d+[)）]', sub_t):
                            has_sub_level = True
                            break
                if has_sub_level:
                    break
            # 只要有子标题结构模式，就提升所有 X.（非X.Y）为 h1
            if has_sub_level and len(x_prefix_paras) >= 2:
                promote_x_to_h1 = True
                # 确定标题区域结束位置（首个 X. 前缀段落之前）
                title_end = min(x_prefix_paras) if x_prefix_paras else 0
                # 收集需要提升的 body 段落：仅限无 Word 编号、无编号前缀、
                # 长度极短（<=15字）且夹在两个 X. 段落之间的纯标题行
                for idx, item in enumerate(paragraphs_text):
                    if idx < title_end:
                        continue  # 跳过标题区域
                    if item[0] == 'p' and item[1].strip():
                        t = item[1].strip()
                        # 有 Word 编号的段落一概不提升
                        is_word_num = (len(item) > 4 and item[4] and item[4] != '0')
                        if is_word_num:
                            continue
                        # 只提升极短的无编号标题（<=15字，排除长句）
                        is_short = (len(t) <= 15
                                    and not re.search(r'[。；]', t)
                                    and not re.search(r'^（\d+）', t)
                                    and not t.startswith('附件'))
                        if not is_short:
                            continue
                        # 检查前后是否有 X. 前缀段落（严格只看 X. 段落，不看 Word 编号段落）
                        has_x_neighbor = False
                        for offset in (-1, -2, 1, 2):
                            ni = idx + offset
                            if 0 <= ni < len(paragraphs_text) and paragraphs_text[ni][0] == 'p':
                                n_t = paragraphs_text[ni][1].strip()
                                if re.match(r'^\d+[.、．]\s*\S', n_t) and not re.match(r'^\d+\.\d+', n_t):
                                    has_x_neighbor = True
                                    break
                        if has_x_neighbor:
                            promote_body_indices.add(idx)

    # 全局规则：主标题上方空一行
    top_blank = doc.add_paragraph()
    set_para_spacing(top_blank)

    title_mode = True
    title_count = 0  # 连续标题段计数，防止将正文标题误判为主标题
    title_ended = False  # 标记标题区是否已结束
    counter = HeadingCounter()
    merged_titles = set()  # 已合并到主标题的段落索引集合

    # 预计算：标记哪些段落索引最终是标题（用于空行过滤）
    is_heading_index = set()

    def _precompute_heading(idx, item):
        """预计算单个段落的最终层级（不含 counter），返回 level"""
        if item[0] != 'p':
            return None
        raw_text = item[1]
        text = clean_text(raw_text)
        if not text:
            return None
        if is_main_title(text):
            return 'title'
        level = detect_level(text)
        wnl = item[3] if len(item) > 3 else None
        # Bug2修复：优先使用detect_level的文本编号模式识别结果
        # 只有detect_level返回body（无法从文本识别标题）时，才用wnl作回退
        # 避免Word编号映射错误覆盖正确的文本识别（如"四、"被wnl='h4'覆盖）
        if wnl is not None and level == 'body':
            level = wnl
        if promote_x_to_h1 and level == 'h3':
            is_long_sentence = '。' in text and len(text) > 30
            if (re.match(r'^\d+[.、．]\s*\S', text)
                and not re.match(r'^\d+\.\d+', text)
                and not re.match(r'^\d+[.、．]\s*[是是以要为将把让使被]', text)
                and not is_long_sentence):
                level = 'h1'
        if promote_x_to_h1 and level == 'body' and wnl is None and idx in promote_body_indices:
            level = 'h1'
        # Word 编号隐藏了编号的短段落（无文本前缀 ≤15字），promote为 h1
        if promote_x_to_h1 and wnl is not None and wnl in ('h1', 'h2', 'h3', 'h4'):
            has_text_prefix = bool(
                re.match(r'^[一二三四五六七八九十]+、', text)
                or re.match(r'^（[一二三四五六七八九十]+）', text)
                or re.match(r'^\d+[.、．]\s*', text)
                or re.match(r'^（\d+）', text)
                or re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', text)
            )
            if not has_text_prefix and len(text) <= 15 and not re.search(r'[。；]', text):
                level = 'h1'
        # Word 编号段落中，文本无编号前缀且长度 >25 字的，不算标题
        # Bug2修复：但如果Word编号格式是标题型格式（chineseCounting等），不降级
        if wnl is not None and level in ('h1', 'h2', 'h3', 'h4', 'h5'):
            has_text_prefix = bool(
                re.match(r'^[一二三四五六七八九十]+、', text)
                or re.match(r'^（[一二三四五六七八九十]+）', text)
                or re.match(r'^\d+[.、．]\s*', text)
                or re.match(r'^（\d+）', text)
                or re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', text)
            )
            if not has_text_prefix and len(text) > 25:
                # 检查Word编号格式
                _orig_nid = item[4] if len(item) > 4 else None
                _orig_nilvl = item[5] if len(item) > 5 else '0'
                _is_heading_num_format = False
                if _orig_nid and _orig_nid != '0':
                    _an_id = num_to_abstract.get(_orig_nid)
                    if _an_id:
                        _levels = abstract_num_defs.get(_an_id, {})
                        _fmt, _ = _levels.get(_orig_nilvl or '0', (None, None))
                        if _fmt in ('chineseCounting', 'chineseCountingThousand',
                                    'ideographDigital', 'ideographEnclosedCircle'):
                            _is_heading_num_format = True
                if not _is_heading_num_format:
                    return None
        return level

    for idx, item in enumerate(paragraphs_text):
        level = _precompute_heading(idx, item)
        if level and level in ('title', 'h1', 'h2', 'h3', 'h4', 'h5'):
            is_heading_index.add(idx)

    # ─── 检查正文编号序号是否合理 ───
    # 规则：同一(numId, ilvl)组内，相邻两段之间如果隔了一个 h1 标题，
    # 说明序号跨大节连续，可能不正确，加入批注警告
    # 注意：h1 标题可能是：
    #  1) 文本前缀"一、二、三、"（如手工录入的文档）
    #  2) Word原生编号(numId=1)的段落（如本模板文档）
    # 两种情况都要检测
    discontinuous_seq_warnings = {}
    # 建立每个(numId, ilvl)组的段落索引列表
    num_group_indices = {}
    for pi, pitem in enumerate(paragraphs_text):
        if pitem[0] != 'p':
            continue
        orig_num_id = pitem[4] if len(pitem) > 4 else None
        orig_num_ilvl = pitem[5] if len(pitem) > 5 else None
        if not orig_num_id or orig_num_id == '0':
            continue
        key = (orig_num_id, orig_num_ilvl)
        if key not in num_group_indices:
            num_group_indices[key] = []
        num_group_indices[key].append(pi)
    # 检查：同一组内，相邻两段之间是否有 h1 标题
    for key, indices in num_group_indices.items():
        for j in range(1, len(indices)):
            prev_idx = indices[j - 1]
            curr_idx = indices[j]
            # 检查 prev_idx 和 curr_idx 之间是否有 h1 标题
            has_h1_between = False
            for k in range(prev_idx + 1, curr_idx):
                if k < len(paragraphs_text) and paragraphs_text[k][0] == 'p':
                    t = paragraphs_text[k][1].strip() if len(paragraphs_text[k]) > 1 else ''
                    # 情况1：文本前缀中文编号（一、二、三、）
                    if re.match(r'^[一二三四五六七八九十]+、', t):
                        has_h1_between = True
                        break
                    # 情况2：Word原生编号 numId=1 的段落（h1标题段落）
                    h1_num_id = paragraphs_text[k][4] if len(paragraphs_text[k]) > 4 else None
                    if h1_num_id == '1':
                        has_h1_between = True
                        break
            if has_h1_between:
                prev_seq = num_seq.get(prev_idx)
                curr_seq = num_seq.get(curr_idx)
                if prev_seq is not None and curr_seq is not None and curr_idx in is_heading_index:
                    # 只有当前段落是标题时才触发序号不连续警告
                    # 正文列表（如"需协同解决的事项"下的1. 2.）不触发此警告
                    discontinuous_seq_warnings[curr_idx] = (
                        f'序号 {curr_seq}. 与前一项（序号 {prev_seq}.）'
                        f'之间隔有大节标题，序号可能不连续，建议确认原文'
                    )

    last_was_h1 = False  # 用于h1间空行
    for i, item in enumerate(paragraphs_text):
        # 跳过已合并到主标题的段落
        if i in merged_titles:
            continue
        
        etype = item[0]
        if etype == 'tbl':
            rows_data = item[1]
            rows_shading = item[2] if len(item) > 2 else None
            rows_font_color = item[3] if len(item) > 3 else None
            rebuild_table(doc, rows_data, rows_shading, rows_font_color)
            continue

        raw = item[1]
        is_bold = item[2] if len(item) > 2 else False
        word_num_level = item[3] if len(item) > 3 else None
        bold_runs = item[6] if len(item) > 6 else []
        # 兜底：若 tuple 只有6元素，从 text 匹配反查 bold_runs_by_elem
        if not bold_runs:
            for ri in bold_runs_by_elem.values():
                total = sum(l for l,_ in ri)
                if total > 0 and total == len(raw.strip()):
                    bold_runs = ri
                    break
        has_mixed_runs = (bold_runs and len(bold_runs) > 1
                          and any(b for _, b in bold_runs)
                          and not all(b for _, b in bold_runs))

        text = clean_text(raw)
        if not text:
            # 空行：如果下一个非空段落是标题，跳过此空行
            skip_empty = False
            for j in range(i + 1, len(paragraphs_text)):
                if paragraphs_text[j][0] == 'p' and paragraphs_text[j][1].strip():
                    if j in is_heading_index:
                        skip_empty = True
                    break
                if paragraphs_text[j][0] == 'tbl':
                    break
            if skip_empty:
                continue
            p = doc.add_paragraph()
            set_para_spacing(p)
            continue

        # 排除带 Word 编号的段落（有 numId 说明是标题，不是主标题的一部分）
        has_word_num = len(item) > 4 and item[4] and item[4] != '0'
        if title_mode and is_main_title(text) and not has_word_num:
            title_count += 1
            title_mode = False  # 只处理第一个主标题，后续不再合并
            
            # 收集连续的主标题段落，合并为完整标题
            # 合并条件：is_main_title 为真，且满足以下之一：
            #   1. 加粗 → 典型主标题碎片
            #   2. 无 Word 编号 + 很短(≤15字) → 如"及数据治理情况汇报"
            # 不合并：有 Word 编号的 → 是标题（如一、xxx），不是主标题碎片
            main_title_parts = [text]
            j = i + 1
            while j < len(paragraphs_text) and paragraphs_text[j][0] == 'p':
                next_text = clean_text(paragraphs_text[j][1])
                next_is_bold = paragraphs_text[j][2] if len(paragraphs_text[j]) > 2 else False
                next_has_wnum = (len(paragraphs_text[j]) > 4
                                 and paragraphs_text[j][4]
                                 and paragraphs_text[j][4] != '0')
                should_merge = (next_is_bold
                                or (not next_has_wnum and len(next_text) <= 15))
                # 排除问候语/称呼（如"尊敬的各位领导："）
                is_greet_text = next_text.startswith('尊敬') or '各位领导' in next_text
                if next_text and is_main_title(next_text) and should_merge and not is_greet_text:
                    # 检查是否紧跟其后（中间无空行或其他内容）
                    if not any(clean_text(paragraphs_text[k][1]) for k in range(i+1, j)):
                        main_title_parts.append(next_text)
                        j += 1
                    else:
                        break
                else:
                    break
            
            # 合并主标题（无缝拼接，让文字自然换行，避免强制断行）
            main_title_parts = main_title_parts[:2]
            combined_title = ''.join(main_title_parts)
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_para_spacing(p)
            run = p.add_run(combined_title)
            set_run_font(run, FONT_XIAOBIAOSONG, SIZE_ERHAO, bold=False)
            
            # 跳过已合并的主标题段落（通过continue外部for循环的逻辑）
            # 记录需要跳过的起始索引，外部循环会处理
            for skip_i in range(i + 1, j):
                merged_titles.add(skip_i)
            
            # 主标题后始终插入空行（与正文区隔）
            title_mode = False
            blank = doc.add_paragraph()
            set_para_spacing(blank)
            continue

        title_mode = False
        level = detect_level(text)
        # 先应用 Word 编号层级（PENDING → 实际层级）
        # 但需谨慎：word_num_level 可能不准确（如"协调解决事项"被误判为h1）
        # 仅当 text 看起来像对应层级的标题时才应用
        if word_num_level is not None and level == 'body':
            # Bug2修复：只有detect_level返回body时才用word_num_level回退
            # 避免Word编号映射错误覆盖正确的文本编号识别
            # 保守策略：仅当文本无编号前缀且较短（≤30字）时才应用word_num_level
            # 如果文本看起来像正文（有句号、逗号、长度>30），保持原level判断
            looks_like_title = (
                len(text) <= 30
                and not re.search(r'[。；]', text)
            )
            if looks_like_title:
                level = word_num_level
            # 否则保持原level（可能是body）
        # 调试：查看应用word_num_level后的最终level
        if 18 <= i <= 22:
            print(f'[DEBUG-FINAL-LEVEL] i={i}, text="{text[:30]}", detect_level={detect_level(text)}, wnl={word_num_level}, final_level={level}, promote={promote_x_to_h1}', flush=True)
            print(f'                   promote_body_indices包含i? {i in promote_body_indices}, word_num_indices包含i? {i in word_num_indices}', flush=True)
        # Word 编号段落中，文本无编号前缀且长度 >25 字的，降级为 body（是正文不是标题）
        # Bug2修复：但如果Word编号格式是chineseCounting（一、二、三）等标题型格式，
        # 说明该段落是Word编号隐藏了标题编号，不应仅因文本长度降级
        if word_num_level is not None and level in ('h1', 'h2', 'h3', 'h4', 'h5'):
            has_text_prefix = bool(
                re.match(r'^[一二三四五六七八九十]+、', text)
                or re.match(r'^（[一二三四五六七八九十]+）', text)
                or re.match(r'^\d+[.、．]\s*', text)
                or re.match(r'^（\d+）', text)
                or re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', text)
            )
            if not has_text_prefix and len(text) > 25:
                # 检查Word编号格式：如果是标题型编号格式，不降级
                _orig_nid = item[4] if len(item) > 4 else None
                _orig_nilvl = item[5] if len(item) > 5 else '0'
                _is_heading_num_format = False
                if _orig_nid and _orig_nid != '0':
                    _an_id = num_to_abstract.get(_orig_nid)
                    if _an_id:
                        _levels = abstract_num_defs.get(_an_id, {})
                        _fmt, _ = _levels.get(_orig_nilvl or '0', (None, None))
                        if _fmt in ('chineseCounting', 'chineseCountingThousand',
                                    'ideographDigital', 'ideographEnclosedCircle'):
                            _is_heading_num_format = True
                if not _is_heading_num_format:
                    level = 'body'
                    word_num_level = None  # 重置，后续不按标题处理
        # 提升 X. → h1（当文档没有中文一级编号时）
        # 排除：X.Y 多级编号、X.是/要 等动词前缀、正文长句
        if promote_x_to_h1 and level == 'h3':
            is_long_sentence = '。' in text and len(text) > 30
            if (re.match(r'^\d+[.、．]\s*\S', text)
                and not re.match(r'^\d+\.\d+', text)
                and not re.match(r'^\d+[.、．]\s*[是是以要为将把让使被]', text)
                and not is_long_sentence):
                level = 'h1'
        # 提升 promote 模式下夹在编号组之间的无编号短段落为 h1
        if promote_x_to_h1 and level == 'body' and word_num_level is None and i in promote_body_indices:
            level = 'h1'
        # promote 模式下：Word 编号隐藏了编号的短段落，可能需要提升为标题
        # 但仅当 word_num_level 是 h2/h3/h4 时才提升（h1已是最高级，不应再提升）
        # 且需要更严格的判断：文本应像标题（不含句号，长度适中）
        if promote_x_to_h1 and word_num_level is not None and word_num_level in ('h2', 'h3', 'h4'):
            has_text_prefix = bool(
                re.match(r'^[一二三四五六七八九十]+、', text)
                or re.match(r'^（[一二三四五六七八九十]+）', text)
                or re.match(r'^\d+[.、．]\s*', text)
                or re.match(r'^（\d+）', text)
                or re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', text)
            )
            # 更严格的标题判断：无句末标点、长度适中、不包含逗号顿号等正文特征
            looks_like_heading = (
                not has_text_prefix
                and len(text) <= 15
                and not re.search(r'[。；！？]', text)  # 无句末标点
                and not re.match(r'^\d+[.、．]', text)  # 无数字编号前缀
            )
            if looks_like_heading:
                level = 'h1'

        if level == 'body' and word_num_level is None:
            prev_etype = paragraphs_text[i - 1][0] if i > 0 else None
            prev_info = paragraphs_text[i - 1] if i > 0 else None
            is_after_table = (prev_etype == 'tbl')
            # 检查前一段是否是加粗（主标题）
            prev_is_bold = (prev_info[0] == 'p' and len(prev_info) > 2 and prev_info[2] is True) if prev_etype == 'p' and prev_info else False
            # 表格标题（表1、表2…、表3-1…）不算标题
            is_tbl_title = is_table_title(text)
            # 日期行（如"（2026年4月28日）"）不算标题
            is_dt_line = is_date_line(text)
            is_short_title = (
                len(text) <= 25
                and not re.search(r'[。；]', text)
                and not text.startswith('附件')
                and not re.match(r'^[\d,.\-+%：:（）()]+$', text)
                and not is_tbl_title
                and not is_dt_line
            )
            # 放宽短标题检测：主标题后面的短文本也应该识别为h1
            # 例如："新开项目上线及业财集成情况" 紧跟主标题，应为"一、新开项目..."
            is_after_main_title = (prev_is_bold and len(text) <= 20)
            if is_bold and is_short_title:
                level = 'h1'
            elif is_after_table and is_short_title and len(text) <= 20 and not is_tbl_title:
                level = 'h1'
            elif is_after_main_title and is_short_title and not is_dt_line:
                # 主标题后面的短文本提升为一级标题（日期行除外）
                level = 'h1'

        clean_heading = text
        std_prefix_match = None
        is_multilevel = bool(re.match(r'^\d+\.\d+', text))

        # 检测 X.是 / X.要 / X.以 等编号+动词的不规范格式
        is_verb_prefix = is_verb_after_number(text)

        if not is_multilevel:
            std_prefix_match = (
                re.match(r'^[一二三四五六七八九十]+、', text)
                or re.match(r'^（[一二三四五六七八九十]+）', text)
                or re.match(r'^\d+[.、．]\s*', text)
                or re.match(r'^（\d+）', text)
            )
        if std_prefix_match and level in ('h1', 'h2', 'h3', 'h4', 'h5'):
            clean_heading = text[std_prefix_match.end():].lstrip()

        # 全局规则：标题原文加粗的，格式化后也保持加粗
        preserve_bold = should_preserve_bold(level, is_bold, text)

        if is_multilevel and level in ('h1', 'h2', 'h3', 'h4', 'h5'):
            # X.Y 多级编号：保留原文编号，不自动重编，添加批注提醒
            p = doc.add_paragraph()
            apply_heading_format(p, level, text, preserve_bold=preserve_bold, bold_runs=bold_runs)
            if i in subhead_para_indices:
                num_m = re.match(r'^(\d+\.\d+[.、．]?)', text)
                num_str = num_m.group(1) if num_m else ''
                comment_list.append((text[:20],
                    f'此处使用"{num_str}"多级编号，建议改为四级标题①②③格式',
                    'multi_level_num'))
        elif is_verb_prefix and level in ('h1', 'h2', 'h3', 'h4', 'h5'):
            # X.是/要 等不规范编号：保留原文编号不重编，添加批注提醒
            verb_m = re.match(r'^(\d+[.、．])', text)
            prefix_text = verb_m.group(1) if verb_m else ''
            p = doc.add_paragraph()
            apply_heading_format(p, 'body', text, preserve_bold=preserve_bold, bold_runs=bold_runs)
            # 加强制加粗：python-docx API 可能被后续覆盖，用 XML 兜底
            if preserve_bold and p.runs:
                if '商务部' in text: print(f'[FORCE-BOLD] running XML fix for: {text[:30]}')
                rPr = p.runs[0]._r.get_or_add_rPr()
                b_e = rPr.find(qn('w:b'))
                if b_e is None:
                    b_e = OxmlElement('w:b')
                    rPr.append(b_e)
                if b_e.get(qn('w:val')) is not None:
                    b_e.set(qn('w:val'), 'true')
            if i in h3_para_indices:
                comment_list.append((text[:20],
                    f'编号"{prefix_text}"后直接跟动词，建议改为"一是…""二是…"格式',
                    'verb_after_num'))
        elif level in ('title', 'h1', 'h2', 'h3', 'h4', 'h5'):
            # 空行规则：h1间、附件前
            if last_was_h1 and (level == 'h1' or level == 'h3' or '附件' in text[:8]):
                _add_blank_line(doc)
            
            # 调试：如果是编号段落，打印详细信息
            if text and re.match(r'^\d+[.、．]', text):
                print(f'[DEBUG-NUM] i={i}, level={level}, counter状态=h1={counter.h1},h2={counter.h2},h3={counter.h3},h4={counter.h4}, text="{text[:40]}"')
            
            # 优先保留原文编号前缀，而不是强制重编
            # 扩展正则，匹配所有标题前缀格式：h1=一、 h2=（一） h3=1. h4=（1） h5=①
            orig_num_match = re.match(
                r'^([一二三四五六七八九十]+[、]|[0-9]+[.、．]|（[一二三四五六七八九十]+）|（[0-9]+）|[①②③④⑤⑥⑦⑧⑨⑩])\s*',
                text
            )
            if orig_num_match:
                # 原文有编号前缀，提取并保留
                orig_prefix = orig_num_match.group(0)
                remaining_text = text[orig_num_match.end():]
                display = orig_prefix + remaining_text
                prefix_used = orig_prefix
                heading_body = remaining_text
                # bold_runs 基于 raw_text，需补偿 clean_text 删除的前导字符
                clean_offset = len(item[1]) - len(text)
                if bold_runs and (clean_offset > 0 or len(orig_prefix.rstrip()) > 0):
                    prefix_len = clean_offset + len(orig_prefix.rstrip())
                    first_rlen, first_rbold = bold_runs[0]
                    new_len = first_rlen - prefix_len
                    if new_len > 0:
                        bold_runs[0] = (new_len, first_rbold)
                    else:
                        bold_runs = bold_runs[1:]  # 前缀完全覆盖，去掉
                print(f'[DEBUG]   -> 保留原文编号="{prefix_used.strip()}", counter保持不变')
                # FIX: 更新counter状态，使其与原文编号一致
                # 解析orig_prefix中的编号值
                num_val = None
                prefix_level = level  # 使用当前detect到的level
                # h1格式："一、" "二、"
                m_h1 = re.match(r'[一二三四五六七八九十]+', orig_prefix)
                if m_h1:
                    num_val = CNUM_TO_INT.get(m_h1.group())
                # h3格式："1." "2."
                if num_val is None:
                    m_h3 = re.match(r'[0-9]+', orig_prefix)
                    if m_h3 and re.match(r'[.、．]', orig_prefix[len(m_h3.group()):]):
                        num_val = int(m_h3.group())
                # h2格式："（一）"
                if num_val is None:
                    m_h2 = re.match(r'（([一二三四五六七八九十]+)）', orig_prefix)
                    if m_h2:
                        num_val = CNUM_TO_INT.get(m_h2.group(1))
                # h4格式："（1）"
                if num_val is None:
                    m_h4 = re.match(r'（([0-9]+)）', orig_prefix)
                    if m_h4:
                        num_val = int(m_h4.group(1))
                # h5格式："①"
                if num_val is None:
                    m_h5 = re.match(r'[①②③④⑤⑥⑦⑧⑨⑩]', orig_prefix)
                    if m_h5:
                        try:
                            num_val = CIRCLE_NUMBERS.index(m_h5.group()) + 1
                        except ValueError:
                            pass
                # 更新counter
                if num_val is not None and prefix_level is not None:
                    if prefix_level == 'h1':
                        counter.h1 = num_val
                    elif prefix_level == 'h2':
                        counter.h2 = num_val
                    elif prefix_level == 'h3':
                        counter.h3 = num_val
                    elif prefix_level == 'h4':
                        counter.h4 = num_val
                    elif prefix_level == 'h5':
                        counter.h5 = num_val
                    print(f'[DEBUG-COUNTER] 更新counter: {prefix_level}={num_val}, 新状态=h1={counter.h1},h2={counter.h2},h3={counter.h3},h4={counter.h4},h5={counter.h5}')
            else:
                # 原文无编号前缀，使用counter生成
                prefix = counter.next(level)
                display = prefix + clean_heading
                prefix_used = prefix
                print(f'[DEBUG] i={i} -> prefix="{prefix}", 新counter状态=h1={counter.h1},h2={counter.h2},h3={counter.h3},h4={counter.h4}')
                heading_body = clean_heading
            
            # 空行规则：h1间、附件前、版记前
            if last_was_h1 and level == 'h1':
                _add_blank_line(doc)
            elif '附件' in text[:8] and last_was_h1:
                _add_blank_line(doc)

            
            p = doc.add_paragraph()
            apply_heading_format(p, level, heading_body, prefix=prefix_used, preserve_bold=preserve_bold, bold_runs=bold_runs)
            if i in punct_para_indices:
                comment_list.append((display[:20],
                    '此标题/段落可能缺少句末标点，请人工确认',
                    'missing_end_punct'))
            elif i in title_punct_para_indices:
                comment_list.append((display[:20],
                    '标题末尾不应有标点符号',
                    'trailing_punct'))
            # 检测阿拉伯数字编号格式（如"1. xxx"），判断是否需要提示改中文格式
            # 公文规范：h1=一、  h2=（一）  h3=1.  h4=（1）
            # 如果文档已按规范使用（一）作为h2，则h3使用"1."是正确格式，不需提示
            # 仅当层级编号格式不匹配规范时才提示
            arabic_num_match = re.match(r'^(\d+)([.、．])', text)
            if arabic_num_match and level in ('h1', 'h2', 'h3'):
                # 判断文档的二级标题格式是否已经是中文（（一）格式）
                # 如果是，则三级标题用"1."是规范写法，不提示
                # 只有在一级标题直接用"1."（无中文一、二、三）时才需要提示
                if level == 'h1':
                    # 一级标题用了"1、"格式，建议改为"一、"
                    num_str = arabic_num_match.group(1)
                    sep = arabic_num_match.group(2)
                    orig_num_prefix = num_str + sep
                    chinese_num = CNUM.get(num_str, num_str)
                    comment_list.append((orig_num_prefix,
                        f'建议将"{orig_num_prefix}"改为"{chinese_num}、"',
                        'number_prefix'))
                elif level == 'h2' and not has_cn_h1:
                    # 二级标题用了"1、"格式，且文档无中文一级标题，建议改为"（一）"
                    num_str = arabic_num_match.group(1)
                    sep = arabic_num_match.group(2)
                    orig_num_prefix = num_str + sep
                    chinese_num = CNUM.get(num_str, num_str)
                    comment_list.append((orig_num_prefix,
                        f'建议将"{orig_num_prefix}"改为"（{chinese_num}）"',
                        'number_prefix'))
            # 检查是否缺少二级标题
            if i in missing_h2_indices:
                comment_list.append((text[:30],
                    '一级标题下直接使用三级标题，建议补充二级标题（一）（二）',
                    'heading_skip'))
            # 检查标题编号标点是否符合规范
            if i in title_punct_indices:
                # 找到对应的问题，获取建议文本
                        for idx, txt, issue_type, suggestion in title_punct_check_results:
                            if idx == i:
                                comment_list.append((text[:30], suggestion, 'number_prefix'))
                                break
            if level == 'h1':
                last_was_h1 = True
        else:
            # 正文段落
            p = doc.add_paragraph()
            # 问候语（含称呼关键词+以：结尾）不缩进
            is_greet = is_greeting(text)
            # 全局规则：表格标题（表1、表2、表3等）居中
            is_tbl_title2 = is_table_title(text)
            
            # 对于有numId的正文段落（Word原生编号列表），保留原始编号格式
            orig_num_id = item[4] if len(item) > 4 else None
            orig_num_ilvl = item[5] if len(item) > 5 else None
            display_text = text
            if orig_num_id and orig_num_id != '0':
                an_id = num_to_abstract.get(orig_num_id)
                if an_id:
                    levels = abstract_num_defs.get(an_id, {})
                    fmt, lvl_txt = levels.get(orig_num_ilvl or '0', (None, None))
                    # lvl_txt 是 Word 的编号模板，如 '（%1）'、'%1.'、'①' 等
                    is_non_decimal = fmt and fmt not in ('decimal', 'none', None, '')
                    is_decimal_no_prefix = (
                        fmt == 'decimal'
                        and not bool(re.match(r'^\d+[.、．]\s*\S', text))
                    )
                    if is_non_decimal:
                        # 非十进制格式（如①②③④），添加编号前缀到文本
                        seq_val = num_seq.get(i, 1)
                        circle_num = CIRCLE_NUMBERS[seq_val - 1] if 1 <= seq_val <= len(CIRCLE_NUMBERS) else str(seq_val)
                        display_text = circle_num + text
                        fmt_name = {
                            'ideographEnclosedCircle': '①②③④',
                            'decimalEnclosedCircleChinese': '①②③④',
                            'chineseCountingThousand': '中文千位',
                            'chineseCounting': '中文计数',
                            'lowerLetter': 'a.b.c.',
                            'upperLetter': 'A.B.C.',
                        }.get(fmt, fmt)
                        if i not in discontinuous_seq_warnings:
                            comment_list.append((circle_num,
                                f'原文使用编号{fmt_name}，是Word自带编号。建议改为（1）（2）或保留原格式',
                                'number_prefix'))
                    elif is_decimal_no_prefix:
                        # 连续编号段落 → 列表项，加前缀（但有混合粗体的不加，粗体已是结构）
                        is_list_item = False
                        if orig_num_id and orig_num_id != '0' and not has_mixed_runs:
                            if i > 0:
                                prev_item = paragraphs_text[i - 1]
                                pnid = prev_item[4] if len(prev_item) > 4 else None
                                if pnid == orig_num_id:
                                    is_list_item = True
                            if not is_list_item and i + 1 < len(paragraphs_text):
                                next_item = paragraphs_text[i + 1]
                                nnid = next_item[4] if len(next_item) > 4 else None
                                if nnid == orig_num_id:
                                    is_list_item = True
                        if is_list_item:
                            seq_val = num_seq.get(i, 1)
                            prefix = lvl_txt.replace('%1', str(seq_val)) if lvl_txt else f'{seq_val}.'
                            display_text = prefix + text
                            # 审查提示：分隔符含顿号
                            if ('、' in prefix or '．' in prefix) and i in sep_indices:
                                good = prefix.replace('、', '.').replace('．', '.')
                                comment_list.append((display_text[:25], f'编号分隔符建议将"{prefix.strip()}"改为"{good}"', 'number_prefix'))
            
            # 逐 run 重建（保留原文局部加粗）
            has_mixed = (bold_runs and len(bold_runs) > 1
                         and any(b for _, b in bold_runs)
                         and not all(b for _, b in bold_runs))
            if has_mixed and display_text == text:
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                set_para_spacing(p)
                if not (is_greet or is_tbl_title2):
                    set_para_indent(p, 2)
                pos = 0
                for rlen, rbold in bold_runs:
                    chunk = text[pos:pos + rlen]
                    if chunk.strip():
                        run = p.add_run(chunk)
                        set_run_font(run, FONT_FANGSONG, SIZE_SANHAO, bold=rbold)
                    pos += rlen
            elif has_mixed_runs and display_text != text:
                # display_text 含编号前缀，逐 run 重建时前缀合并到第一个 run
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                set_para_spacing(p)
                if not (is_greet or is_tbl_title2):
                    set_para_indent(p, 2)
                prefix = display_text[:len(display_text) - len(text)]
                first_rlen, first_rbold = bold_runs[0]
                first_chunk = text[:first_rlen]
                run0 = p.add_run(prefix + first_chunk)
                set_run_font(run0, FONT_FANGSONG, SIZE_SANHAO, bold=first_rbold)
                pos = first_rlen
                for rlen, rbold in bold_runs[1:]:
                    chunk = text[pos:pos + rlen]
                    if chunk.strip():
                        run = p.add_run(chunk)
                        set_run_font(run, FONT_FANGSONG, SIZE_SANHAO, bold=rbold)
                    pos += rlen
            else:
                apply_heading_format(p, level, display_text, no_indent=is_greet or is_tbl_title2, preserve_bold=preserve_bold, bold_runs=bold_runs)
            # 表格标题居中且不加粗（保持原文加粗状态）
            if is_tbl_title2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i in punct_para_indices:
                comment_list.append((text[:20],
                    '此段落可能缺少句末标点，请人工确认',
                    'missing_end_punct'))
            if i in discontinuous_seq_warnings:
                # 用 text[:30] 而不是 display_text[:30]，因为输出文档的纯文本不含编号前缀
                comment_list.append((text[:30],
                    discontinuous_seq_warnings[i],
                    'number_prefix'))
            # Bug 1修复：Word自动编号格式批注
            if i in word_num_indices:
                comment_list.append((text[:30],
                    '原文使用Word自动编号（阿拉伯数字1.2.3.），建议改为二级标题（（一）（二））格式',
                    'number_prefix'))

            last_was_h1 = False

    # 注入编号序号问题的批注
    comment_list.extend(list_num_issues)

    # ──── 全局规则：落款格式处理 ────
    # 署名单位右对齐 + 日期右空4字 + 与上文空两行
    # 规则：落款只在全文末尾，不在文档开头或中间
    doc_paras = list(doc.paragraphs)
    total_paras = len(doc_paras)
    # 忽略末尾空行，计算有效末尾位置
    last_content_idx = total_paras - 1
    while last_content_idx >= 0 and not doc_paras[last_content_idx].text.strip():
        last_content_idx -= 1
    for pi in range(total_paras - 1, 0, -1):
        if is_signature_date(doc_paras[pi].text):
            # 只处理靠近末尾的日期（忽略尾随空行）
            if last_content_idx - pi > 5:
                continue
            # 日期后不能有正文（标题或长段落）
            has_content_after = False
            for ck in range(pi + 1, total_paras):
                ct = doc_paras[ck].text.strip()
                if not ct:
                    continue
                if len(ct) > 15 or re.match(r'^[一二三四五六七八九十]+、', ct):
                    has_content_after = True
                    break
            if has_content_after:
                continue
            # 日期行：RIGHT + 右空4字（三号16pt × 4 = 64pt）
            dpPr = doc_paras[pi]._p.get_or_add_pPr()
            # 清除旧对齐
            old_jc = dpPr.find(qn('w:jc'))
            if old_jc is not None: dpPr.remove(old_jc)
            jc_date = OxmlElement('w:jc'); jc_date.set(qn('w:val'), 'right')
            dpPr.append(jc_date)
            # 右缩进
            ind_d = dpPr.find(qn('w:ind'))
            if ind_d is None:
                ind_d = OxmlElement('w:ind'); dpPr.append(ind_d)
            ind_d.set(qn('w:right'), str(64 * 20))  # 64pt = 1280 twips（w:ind 用 twips）
            # 署名单位：CENTER + 同右缩进
            if pi > 0:
                spPr = doc_paras[pi - 1]._p.get_or_add_pPr()
                old_jc2 = spPr.find(qn('w:jc'))
                if old_jc2 is not None: spPr.remove(old_jc2)
                jc_sig = OxmlElement('w:jc'); jc_sig.set(qn('w:val'), 'right')
                spPr.append(jc_sig)
                ind_s = spPr.find(qn('w:ind'))
                if ind_s is None:
                    ind_s = OxmlElement('w:ind'); spPr.append(ind_s)
                ind_s.set(qn('w:right'), str(64 * 20))  # 64pt = 1280 twips
                # 在署名单位前插入两个空行
                co_elem = doc_paras[pi - 1]._element
                parent = co_elem.getparent()
                ins_idx = list(parent).index(co_elem)
                for _ in range(2):
                    bp = OxmlElement('w:p')
                    bpPr = OxmlElement('w:pPr')
                    bspc = OxmlElement('w:spacing')
                    bspc.set(qn('w:line'), str(LINE_SPACING_TWIPS))
                    bspc.set(qn('w:lineRule'), 'exact')
                    bpPr.append(bspc)
                    bp.append(bpPr)
                    parent.insert(ins_idx, bp)
            break

    # 统一添加批注
    _apply_comments_to_doc(doc, comment_list)

    _add_page_number(doc)
    doc.save(dst_path)
    return dst_path, warnings


def _add_blank_line(doc):
    """插入一个 28.9pt 行高的空行"""
    p = doc.add_paragraph()
    set_para_spacing(p)

