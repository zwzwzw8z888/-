#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
中建四局公文格式化工具 - Web服务端
功能：格式化 / 内容审查 / 公文起草
"""

import os
import re
import uuid
import json
import zipfile
import threading
import time
from pathlib import Path
from flask import Flask, request, jsonify, send_file, after_this_request
from flask_cors import CORS
from werkzeug.utils import secure_filename

from formatter_core import format_document

app = Flask(__name__, static_folder='.', static_url_path='')
CORS(app)


@app.route('/')
def index():
    return send_file(str(BASE_DIR / 'index.html'))

# ── 目录配置 ──
BASE_DIR   = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "outputs"
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

ALLOWED_EXT = {'.docx', '.txt', '.md'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


# ══════════════════════════════════════════════
#  1. 格式化接口
# ══════════════════════════════════════════════

@app.route('/api/format', methods=['POST'])
def api_format():
    """上传文件 → 返回格式化后的 docx"""
    if 'file' not in request.files:
        return jsonify({'ok': False, 'msg': '未上传文件'}), 400

    f = request.files['file']
    if not f.filename:
        return jsonify({'ok': False, 'msg': '文件名为空'}), 400

    ext = Path(f.filename).suffix.lower()
    if ext not in ALLOWED_EXT:
        return jsonify({'ok': False, 'msg': f'不支持的格式 {ext}，请上传 .docx .txt .md'}), 400

    uid = uuid.uuid4().hex
    src_path = UPLOAD_DIR / f"{uid}{ext}"
    stem = Path(f.filename).stem
    dst_path = OUTPUT_DIR / f"{uid}_{stem}_公文格式.docx"

    try:
        f.save(str(src_path))
        format_document(str(src_path), str(dst_path))
    except Exception as e:
        return jsonify({'ok': False, 'msg': f'转换失败：{e}'}), 500
    finally:
        src_path.unlink(missing_ok=True)

    # 异步清理输出文件（5分钟后）
    def _cleanup():
        time.sleep(300)
        dst_path.unlink(missing_ok=True)
    threading.Thread(target=_cleanup, daemon=True).start()

    return send_file(
        str(dst_path),
        as_attachment=True,
        download_name=f"{stem}_公文格式.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# ══════════════════════════════════════════════
#  1.5 批量格式化接口
# ══════════════════════════════════════════════

@app.route('/api/batch', methods=['POST'])
def api_batch():
    """上传多个文件 → 格式化后打包为 zip 返回"""
    files = request.files.getlist('files')
    if not files:
        return jsonify({'ok': False, 'msg': '未上传文件'}), 400

    uid = uuid.uuid4().hex
    batch_dir = OUTPUT_DIR / f"batch_{uid}"
    batch_dir.mkdir(exist_ok=True)
    zip_path = OUTPUT_DIR / f"batch_{uid}.zip"
    results = []

    for f in files:
        if not f.filename:
            continue
        ext = Path(f.filename).suffix.lower()
        if ext not in ALLOWED_EXT:
            results.append({'file': f.filename, 'ok': False, 'msg': f'不支持格式 {ext}'})
            continue
        if f.content_length and f.content_length > MAX_FILE_SIZE:
            results.append({'file': f.filename, 'ok': False, 'msg': '文件超过10MB'})
            continue

        stem = Path(f.filename).stem
        src_path = UPLOAD_DIR / f"{uid}_{stem}{ext}"
        dst_path = batch_dir / f"{stem}_公文格式.docx"
        try:
            f.save(str(src_path))
            format_document(str(src_path), str(dst_path))
            results.append({'file': f.filename, 'ok': True})
        except Exception as e:
            results.append({'file': f.filename, 'ok': False, 'msg': str(e)[:100]})
        finally:
            src_path.unlink(missing_ok=True)

    # 打包 zip
    with zipfile.ZipFile(str(zip_path), 'w', zipfile.ZIP_DEFLATED) as zf:
        for f in sorted(batch_dir.iterdir()):
            if f.suffix == '.docx':
                zf.write(str(f), f.name)
    # 清理临时目录
    import shutil
    shutil.rmtree(str(batch_dir), ignore_errors=True)
    # 5分钟后清理zip
    def _cleanup_zip():
        time.sleep(300)
        zip_path.unlink(missing_ok=True)
    threading.Thread(target=_cleanup_zip, daemon=True).start()

    return send_file(
        str(zip_path),
        as_attachment=True,
        download_name=f"格式化结果_{len(results)}个文件.zip",
        mimetype='application/zip'
    )


# ══════════════════════════════════════════════
#  2. 内容审查接口
# ══════════════════════════════════════════════

REVIEW_RULES = [
    # (pattern, issue_type, description, suggestion)
    (r'[^\u4e00-\u9fa5\u0000-\u007f\s]', '特殊字符', '包含非常规字符', '请确认是否为误输入'),
    (r'\d{4}年\d{1,2}月\d{1,2}', None, None, None),   # 合法日期，skip
    # 数字使用
    (r'[零一二三四五六七八九十百千万亿]+\s*个|[零一二三四五六七八九十]+\s*项',
     '数字规范', '正文数量应使用阿拉伯数字', '请将中文数字改为阿拉伯数字（日期、序号除外）'),
    # 口语化
    (r'很多|很大|很好|很快|大概|也许|可能|应该|好像|感觉',
     '用词规范', '存在口语化或模糊性词汇', '请使用正式书面表达，如"大量""较多""积极推进"等'),
    # 网络用语
    (r'超级|非常非常|棒棒|哈哈|好吧|嗯嗯',
     '用词规范', '存在非正式表达', '请替换为正式公文用语'),
    # 标点
    (r'！！|？？|……{3,}',
     '标点规范', '存在重复或过多标点符号', '公文应使用规范标点，避免感叹号重复等'),
    # 长句（超过100字无标点）
    (r'[^。！？；\n]{100,}',
     '句式规范', '存在超长无断句语句（超过100字）', '建议适当拆分，保持语句简洁'),
]

TITLE_RE = {
    'h1': re.compile(r'^[一二三四五六七八九十]+、'),
    'h2': re.compile(r'^（[一二三四五六七八九十]+）'),
    'h3': re.compile(r'^\d+[.、．]\s*'),
    'h4': re.compile(r'^（\d+）'),
    'h5': re.compile(r'^[①②③④⑤⑥⑦⑧⑨⑩]'),
}

HEADING_ORDER = ['h1', 'h2', 'h3', 'h4', 'h5']

ENDINGS = {'请示': '妥否，请示', '报告': '特此报告', '批复': '特此批复'}


def _extract_text_from_docx(path: str) -> list[dict]:
    """从 docx 提取段落列表，返回 [{'text': ..., 'bold': ..., 'line': ...}]"""
    from docx import Document
    doc = Document(path)
    result = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        is_bold = any(run.bold for run in para.runs if run.text.strip())
        result.append({'text': text, 'bold': is_bold, 'line': i + 1})
    return result


def _detect_doc_type(paragraphs: list[dict]) -> str:
    """简单判断文种"""
    full_text = ' '.join(p['text'] for p in paragraphs[:5])
    if '请示' in full_text:
        return '请示'
    elif '报告' in full_text:
        return '报告'
    elif '通知' in full_text:
        return '通知'
    elif '批复' in full_text:
        return '批复'
    return '未知'


def review_content(path: str) -> dict:
    """对文档内容进行审查，返回审查结果"""
    issues = []
    stats = {'paragraphs': 0, 'issues': 0, 'headings': {'h1': 0, 'h2': 0, 'h3': 0}}

    try:
        paras = _extract_text_from_docx(path)
    except Exception as e:
        return {'ok': False, 'msg': f'读取文件失败：{e}'}

    stats['paragraphs'] = len(paras)
    doc_type = _detect_doc_type(paras)

    # 标题层级顺序检查
    last_level = None
    for p in paras:
        text = p['text']
        current_level = None
        for lvl, pat in TITLE_RE.items():
            if pat.match(text):
                current_level = lvl
                stats['headings'][lvl] = stats['headings'].get(lvl, 0) + 1
                break
        if current_level:
            if last_level:
                curr_idx = HEADING_ORDER.index(current_level)
                last_idx = HEADING_ORDER.index(last_level)
                if curr_idx > last_idx + 1:
                    issues.append({
                        'level': 'warning',
                        'type': '标题层级',
                        'line': p['line'],
                        'text': text[:30],
                        'msg': f'标题层级跳跃：{last_level} → {current_level}，建议补充中间层级',
                    })
            last_level = current_level

        # 内容规范检查
        for pattern, issue_type, desc, suggestion in REVIEW_RULES:
            if issue_type is None:
                continue
            if re.search(pattern, text):
                issues.append({
                    'level': 'info',
                    'type': issue_type,
                    'line': p['line'],
                    'text': text[:40] + ('…' if len(text) > 40 else ''),
                    'msg': desc,
                    'suggestion': suggestion,
                })
                break

    # 结束语检查
    if doc_type in ENDINGS and paras:
        last_text = paras[-1]['text']
        expected = ENDINGS[doc_type]
        if expected not in last_text:
            issues.append({
                'level': 'error',
                'type': '结束语',
                'line': paras[-1]['line'],
                'text': last_text[:30],
                'msg': f'{doc_type}结束语应为"{expected}"，当前末段未包含',
                'suggestion': f'请在正文末段添加"{expected}"',
            })

    # 标题数量检查
    if stats['headings'].get('h1', 0) == 0:
        issues.append({
            'level': 'warning',
            'type': '结构完整性',
            'line': 0,
            'text': '',
            'msg': '文档缺少一级标题（一、二、三、……）',
            'suggestion': '正式公文正文应有清晰的一级标题划分',
        })

    stats['issues'] = len(issues)
    return {
        'ok': True,
        'doc_type': doc_type,
        'stats': stats,
        'issues': issues,
        'summary': f'共审查 {stats["paragraphs"]} 段，发现 {len(issues)} 条建议',
    }


@app.route('/api/review', methods=['POST'])
def api_review():
    """上传文件 → 返回审查报告（JSON）"""
    if 'file' not in request.files:
        return jsonify({'ok': False, 'msg': '未上传文件'}), 400

    f = request.files['file']
    ext = Path(f.filename).suffix.lower()
    if ext != '.docx':
        return jsonify({'ok': False, 'msg': '内容审查仅支持 .docx 文件'}), 400

    uid = uuid.uuid4().hex
    src_path = UPLOAD_DIR / f"{uid}.docx"
    f.save(str(src_path))

    try:
        result = review_content(str(src_path))
    finally:
        src_path.unlink(missing_ok=True)

    return jsonify(result)


# ══════════════════════════════════════════════
#  3. 公文起草接口
# ══════════════════════════════════════════════

DRAFT_TEMPLATES = {
    '请示': {
        'name': '请示',
        'ending': '妥否，请示。',
        'sections': ['背景与目的', '主要内容', '拟采取的措施', '请示事项'],
    },
    '报告': {
        'name': '报告',
        'ending': '特此报告。',
        'sections': ['工作概况', '主要工作进展', '存在问题', '下一步计划'],
    },
    '通知': {
        'name': '通知',
        'ending': '',
        'sections': ['通知事项', '具体要求', '注意事项'],
    },
    '批复': {
        'name': '批复',
        'ending': '特此批复。',
        'sections': ['批复意见', '具体要求'],
    },
    '函': {
        'name': '函',
        'ending': '特此函告，请知悉。',
        'sections': ['事由说明', '主要内容', '请求/告知事项'],
    },
    '会议纪要': {
        'name': '会议纪要',
        'ending': '',
        'sections': ['会议基本情况', '会议研究事项', '议定事项', '其他事项'],
    },
}


def draft_document(
    doc_type: str,
    title: str,
    sender: str,
    recipient: str,
    date: str,
    content_hints: dict,
    include_header: bool = False,
    doc_number: str = ''
) -> str:
    """
    生成公文起草模板 docx，返回输出路径
    content_hints: {section_name: hint_text}
    """
    from docx import Document as DocX
    from formatter_core import (
        set_run_font, set_para_spacing, set_para_indent,
        _add_page_number,
        FONT_FANGSONG, FONT_HEITI, FONT_KAITI, FONT_XIAOBIAOSONG,
        SIZE_CHUHAO, SIZE_ERHAO, SIZE_SANHAO, SIZE_XIAOSI,
        MARGIN_TOP, MARGIN_BOTTOM, MARGIN_LEFT, MARGIN_RIGHT,
        LINE_SPACING_TWIPS, CN_NUMBERS
    )
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tmpl = DRAFT_TEMPLATES.get(doc_type, DRAFT_TEMPLATES['报告'])

    doc = DocX()
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin   = MARGIN_LEFT
    section.right_margin  = MARGIN_RIGHT

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = SIZE_SANHAO

    def add_para(text, cn_font, size, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 indent=0, color=None, twips=LINE_SPACING_TWIPS):
        p = doc.add_paragraph()
        p.alignment = align
        set_para_spacing(p, twips)
        if indent:
            set_para_indent(p, indent)
        run = p.add_run(text)
        set_run_font(run, cn_font, size, bold=bold, color=color)
        return p

    # ── 红头（可选）
    if include_header:
        add_para(
            sender or '中国建筑第四工程局有限公司',
            FONT_XIAOBIAOSONG, SIZE_CHUHAO,
            color=RGBColor(0xC0, 0x00, 0x00),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            twips=900
        )
        # 分割线
        sep = doc.add_paragraph()
        pPr = sep._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'C00000')
        pBdr.append(bottom)
        pPr.append(pBdr)
        set_para_spacing(sep, twips=240)

    # ── 发文字号
    if doc_number:
        add_para(doc_number, FONT_FANGSONG, SIZE_SANHAO,
                 align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── 标题
    add_para(title or f'关于……的{tmpl["name"]}',
             FONT_XIAOBIAOSONG, SIZE_ERHAO,
             align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── 主送单位
    if recipient:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_para_spacing(p)
        run = p.add_run(recipient + '：')
        set_run_font(run, FONT_FANGSONG, SIZE_SANHAO)

    # ── 正文各节
    for idx, section_name in enumerate(tmpl['sections'], 1):
        cn_idx = CN_NUMBERS[idx - 1] if idx <= len(CN_NUMBERS) else str(idx)
        heading_text = f'{cn_idx}、{section_name}'
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_para_spacing(p)
        set_para_indent(p, 2)
        run = p.add_run(heading_text)
        set_run_font(run, FONT_HEITI, SIZE_SANHAO)

        # 节内容（来自 hints 或默认占位）
        hint = content_hints.get(section_name, f'（请填写{section_name}相关内容）')
        add_para(hint, FONT_FANGSONG, SIZE_SANHAO, indent=2)

    # ── 结束语
    if tmpl['ending']:
        add_para(tmpl['ending'], FONT_FANGSONG, SIZE_SANHAO, indent=2)

    # ── 落款
    doc.add_paragraph()  # 空行
    add_para(sender or '（发文机关）', FONT_FANGSONG, SIZE_SANHAO,
             align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_para(date or '2026年  月  日', FONT_FANGSONG, SIZE_SANHAO,
             align=WD_ALIGN_PARAGRAPH.RIGHT)

    _add_page_number(doc)

    uid = uuid.uuid4().hex
    out_path = OUTPUT_DIR / f"{uid}_{doc_type}草稿.docx"
    doc.save(str(out_path))

    # 5分钟后清理
    def _cleanup():
        time.sleep(300)
        out_path.unlink(missing_ok=True)
    threading.Thread(target=_cleanup, daemon=True).start()

    return str(out_path)


@app.route('/api/draft', methods=['POST'])
def api_draft():
    """起草公文 → 返回 docx"""
    data = request.get_json(silent=True) or {}
    doc_type       = data.get('doc_type', '报告')
    title          = data.get('title', '')
    sender         = data.get('sender', '')
    recipient      = data.get('recipient', '')
    date           = data.get('date', '')
    content_hints  = data.get('content_hints', {})
    include_header = data.get('include_header', False)
    doc_number     = data.get('doc_number', '')

    if doc_type not in DRAFT_TEMPLATES:
        return jsonify({'ok': False, 'msg': f'不支持的文种：{doc_type}'}), 400

    try:
        out_path = draft_document(
            doc_type, title, sender, recipient,
            date, content_hints, include_header, doc_number
        )
    except Exception as e:
        return jsonify({'ok': False, 'msg': f'起草失败：{e}'}), 500

    stem = title or f'{doc_type}草稿'
    return send_file(
        out_path,
        as_attachment=True,
        download_name=f"{stem}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/api/draft_types', methods=['GET'])
def api_draft_types():
    """返回支持的文种列表"""
    return jsonify({
        'ok': True,
        'types': [
            {'key': k, 'name': v['name'], 'sections': v['sections']}
            for k, v in DRAFT_TEMPLATES.items()
        ]
    })


# ── 健康检查
@app.route('/api/health', methods=['GET'])
def api_health():
    return jsonify({'ok': True, 'msg': '服务正常', 'version': '1.0.0'})


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('DEBUG', '0') == '1'
    print(f"🚀 中建四局公文格式化工具启动中 → http://localhost:{port}")
    app.run(host='0.0.0.0', port=port, debug=debug)
